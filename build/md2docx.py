import re, sys, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1B,0x3A,0x5C); RED=RGBColor(0xC1,0x35,0x2B); INK=RGBColor(0x1F,0x29,0x33)
JP='Meiryo'

def setfont(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.size=Pt(size); run.bold=bold; run.italic=italic; run.font.color.rgb=color
    run.font.name=JP
    r=run._element.rPr.rFonts
    r.set(qn('w:eastAsia'),JP); r.set(qn('w:ascii'),JP); r.set(qn('w:hAnsi'),JP)

def shade(cell,hexcolor):
    tcPr=cell._tc.get_or_add_tcPr(); sh=OxmlElement('w:shd')
    sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor); tcPr.append(sh)

INLINE=re.compile(r'(\*\*.+?\*\*)')
def add_runs(par, text, size=10.5, color=INK, base_bold=False):
    if text.count('**') % 2:            # 行をまたぐ強調は行全体を太字にする
        text = text.replace('**',''); base_bold = True
    for part in INLINE.split(text):
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            setfont(par.add_run(part[2:-2]), size, True, color)
        else:
            setfont(par.add_run(part.replace('\\|','|')), size, base_bold, color)

def split_row(line):
    return [c.strip() for c in line.strip().strip('|').split('|')]

def convert(md_path, out_path):
    doc=Document()
    sec=doc.sections[0]
    sec.top_margin=sec.bottom_margin=Cm(2.0); sec.left_margin=sec.right_margin=Cm(2.0)
    st=doc.styles['Normal']; st.font.name=JP; st.font.size=Pt(10.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'),JP)

    lines=open(md_path,encoding='utf-8').read().split('\n')
    i=0
    while i < len(lines):
        ln=lines[i]; s=ln.strip()
        if not s:
            i+=1; continue
        # table
        if s.startswith('|') and i+1<len(lines) and re.match(r'^\|[\s:\-|]+\|$', lines[i+1].strip()):
            hdr=split_row(s); rows=[]; i+=2
            while i<len(lines) and lines[i].strip().startswith('|'):
                rows.append(split_row(lines[i].strip())); i+=1
            t=doc.add_table(rows=1+len(rows), cols=len(hdr)); t.style='Table Grid'
            t.alignment=WD_TABLE_ALIGNMENT.CENTER
            for c,txt in enumerate(hdr):
                cell=t.cell(0,c); cell.text=''; shade(cell,'1B3A5C')
                add_runs(cell.paragraphs[0], txt.replace('**',''), 9.5, RGBColor(0xFF,0xFF,0xFF), True)
            for r,row in enumerate(rows,1):
                for c in range(len(hdr)):
                    cell=t.cell(r,c); cell.text=''
                    if r%2==0: shade(cell,'F1F4F7')
                    add_runs(cell.paragraphs[0], row[c] if c<len(row) else '', 9.5)
            doc.add_paragraph()
            continue
        if s.startswith('```'):
            i+=1; buf=[]
            while i<len(lines) and not lines[i].strip().startswith('```'):
                buf.append(lines[i]); i+=1
            i+=1
            par=doc.add_paragraph(); run=par.add_run('\n'.join(buf))
            run.font.name='Courier New'; run.font.size=Pt(9)
            continue
        if s.startswith('---'):
            i+=1; continue
        if s.startswith('#'):
            lvl=len(s)-len(s.lstrip('#')); txt=s.lstrip('#').strip()
            par=doc.add_paragraph()
            sizes={1:20,2:15,3:12.5,4:11}
            par.paragraph_format.space_before=Pt(14 if lvl<=2 else 10)
            par.paragraph_format.space_after=Pt(6)
            if lvl==1: par.alignment=WD_ALIGN_PARAGRAPH.CENTER
            add_runs(par, txt, sizes.get(lvl,11), NAVY if lvl<=2 else RED, True)
            i+=1; continue
        if s=='>':
            i+=1; continue
        if s.startswith('> '):
            par=doc.add_paragraph(); par.paragraph_format.left_indent=Cm(0.8)
            add_runs(par, s[2:], 10.5, NAVY); i+=1; continue
        m=re.match(r'^(\s*)[-*] (\[[ x]\] )?(.*)$', ln)
        if m:
            par=doc.add_paragraph(style='List Bullet')
            par.paragraph_format.space_after=Pt(2)
            if m.group(1): par.paragraph_format.left_indent=Cm(1.4)
            add_runs(par, ('□ ' if m.group(2) else '')+m.group(3)); i+=1; continue
        m=re.match(r'^(\s*)(\d+)\. (.*)$', ln)
        if m:
            par=doc.add_paragraph(style='List Number')
            par.paragraph_format.space_after=Pt(2)
            add_runs(par, m.group(3)); i+=1; continue
        par=doc.add_paragraph(); par.paragraph_format.space_after=Pt(6)
        add_runs(par, s); i+=1
    doc.save(out_path)

if __name__=='__main__':
    convert(sys.argv[1], sys.argv[2])
